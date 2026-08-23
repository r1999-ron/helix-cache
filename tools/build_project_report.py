from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "HelixCache_Project_Report.docx"
NAVY = "123B4A"
TEAL = "16846B"
MINT = "DDF4EC"
PALE = "EEF7F4"
GRAY = "5D6B70"
LIGHT = "F2F4F7"
WHITE = "FFFFFF"
INK = "172529"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size = "Calibri", Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 16, 8), ("Heading 2", 13, TEAL, 12, 6), ("Heading 3", 12, NAVY, 8, 4)
]:
    s = styles[name]
    s.font.name, s.font.size, s.font.bold = "Calibri", Pt(size), True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ["List Bullet", "List Number"]:
    s = styles[name]
    s.font.name, s.font.size = "Calibri", Pt(11)
    s.paragraph_format.left_indent = Inches(.5)
    s.paragraph_format.first_line_indent = Inches(-.25)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.line_spacing = 1.167

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tc_pr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None: tc_mar.append(node)
    if tc_mar.getparent() is None: tc_pr.append(tc_mar)

def fixed_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360"); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd"); tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa"); table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for i, (cell, text, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        shade(cell, NAVY); margins(cell); cell.width = Inches(width / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), size=10, color=WHITE, bold=True)
        cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(width)); cell._tc.get_or_add_tcPr().tcW.set(qn("w:type"), "dxa")
    for r_i, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, text, width in zip(cells, values, widths):
            margins(cell); cell.width = Inches(width / 1440)
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(width)); cell._tc.get_or_add_tcPr().tcW.set(qn("w:type"), "dxa")
            if r_i % 2: shade(cell, PALE)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(text)), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def callout(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(8), Pt(10)
    p.paragraph_format.left_indent, p.paragraph_format.right_indent = Inches(.18), Inches(.18)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), MINT); pPr.append(shd)
    set_font(p.add_run(label + "  "), size=10.5, color=TEAL, bold=True)
    set_font(p.add_run(text), size=10.5, color=INK)

def flow(lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(10)
    p.paragraph_format.line_spacing = 1.25
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), PALE); p._p.get_or_add_pPr().append(shd)
    set_font(p.add_run("\n".join(lines)), name="Consolas", size=9.5, color=NAVY, bold=True)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

# Running furniture
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("HELIXCACHE  |  PROJECT REPORT"), size=8.5, color=GRAY, bold=True)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Agentic hierarchical memory with simulated DNA archival storage  |  "), size=8, color=GRAY)
field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(78)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("TECHNICAL PROJECT REPORT"), size=10, color=TEAL, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("HelixCache"), size=34, color=NAVY, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(24)
set_font(p.add_run("Agentic Hierarchical Memory for AI Using DNA Archival Storage"), size=15, color=TEAL)
flow(["PREDICT DEMAND", "↓", "GPU  →  RAM  →  SSD  →  S3  →  DNA", "↑", "RESTORE BEFORE INFERENCE"])
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(52)
set_font(p.add_run("A working portfolio experiment in storage optimization, error recovery, and predictive prefetching"), size=11, color=GRAY, italic=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(60)
set_font(p.add_run(f"Prepared August 2026  •  MVP Version 0.1"), size=10, color=GRAY)
doc.add_page_break()

doc.add_heading("Executive summary", level=1)
doc.add_paragraph("HelixCache is a working proof of concept for managing AI artifacts across storage tiers with different speed and cost characteristics. It decides where an artifact should live, represents cold artifacts as DNA sequences, repairs isolated nucleotide substitutions, verifies restored files, and predicts which archived artifacts a future request may need.")
callout("Core result", "A real binary file can be uploaded, compressed, DNA-encoded, damaged in a controlled experiment, recovered through redundant strands, verified by SHA-256, restored, and downloaded byte-for-byte intact.")
doc.add_heading("What I built", level=2)
for item in [
    "A zero-dependency Node.js service with a browser dashboard and JSON API.",
    "Five logical tiers: GPU, RAM, SSD, object storage (S3), and DNA archive.",
    "A gzip and two-bit DNA codec with FASTA-like strand storage.",
    "Triple-replicated strands, majority-vote correction, and checksum verification.",
    "An explainable placement score based on frequency, recency, demand, priority, and size.",
    "Keyword-based dependency prediction, cold-artifact prefetching, and a latency comparison.",
    "Hands-on demonstrations for real-file restoration, DNA corruption, and prefetch performance."
]: bullet(item)

doc.add_heading("1. Problem and motivation", level=1)
doc.add_paragraph("Agentic AI systems accumulate many artifacts: model versions, LoRA adapters, retrieval indexes, documents, datasets, and evaluation history. Keeping every artifact in GPU or RAM is expensive and unnecessary. Moving everything to cold storage saves money but increases waiting time when an old artifact becomes relevant again.")
doc.add_paragraph("HelixCache investigates the control layer between AI applications and heterogeneous storage. Its central idea is not that DNA makes inference faster. The idea is that an intelligent controller can predict slow retrieval early enough to reduce its effect on the user-facing request.")
fixed_table(["Tier", "Simple analogy", "Typical purpose", "Demo implementation"], [
    ("GPU", "Worktable", "Immediate inference", "Local gpu/ folder"),
    ("RAM", "Nearby drawer", "Very frequent artifacts", "Local ram/ folder"),
    ("SSD", "Office cabinet", "Prepared or occasional artifacts", "Local ssd/ folder"),
    ("S3", "Remote warehouse", "Rare artifacts", "Local s3/ folder"),
    ("DNA", "Long-term vault", "Very cold archive", "FASTA-like .dna files"),
], [900, 1700, 3100, 3660])

doc.add_heading("2. System architecture", level=1)
flow(["USER REQUEST", "↓", "REQUEST RESOLVER  →  DEPENDENCY PREDICTION", "↓", "ARTIFACT REGISTRY  →  TIER LOOKUP", "↓", "DNA RETRIEVAL + CHECKSUM VERIFICATION", "↓", "SSD PREFETCH  →  GPU LOAD  →  INFERENCE-READY", "↓", "USAGE UPDATE  →  TIERING POLICY"])
doc.add_heading("Main components", level=2)
fixed_table(["Component", "Responsibility"], [
    ("Dashboard", "Runs the three experiments and displays tiers, scores, events, and benchmark results."),
    ("HTTP API", "Registers artifacts and exposes archive, retrieval, experiment, download, optimization, and benchmark operations."),
    ("Artifact registry", "Stores identity, location, size, checksum, usage, demand, and importance metadata."),
    ("DNA codec", "Compresses bytes, maps two-bit values to bases, creates replicated strands, decodes consensus, and verifies integrity."),
    ("Tiering controller", "Computes placement scores and moves artifacts between tier directories."),
    ("Request resolver", "Matches request terms to artifact IDs and ranks likely dependencies."),
    ("Prefetch benchmark", "Compares sequential retrieval with retrieval overlapped with the planning stage."),
], [2300, 7060])

doc.add_heading("3. DNA archive experiment", level=1)
doc.add_heading("Encoding pipeline", level=2)
flow(["ORIGINAL FILE", "↓ gzip compression", "COMPRESSED BYTES", "↓ 2-bit mapping: 00=A, 01=C, 10=G, 11=T", "DNA BASES", "↓ split + replicate", "FASTA-LIKE DNA ARCHIVE"])
doc.add_paragraph("Each DNA fragment is stored three times. The archive header records codec version, original and compressed sizes, checksum, strand length, copy count, and fragment count. The approach is intentionally simple and inspectable rather than biologically production-ready.")
doc.add_heading("Corruption and recovery", level=2)
doc.add_paragraph("The damage lab changes selected nucleotides in replicated strands. During decoding, HelixCache compares the copies at each base position and selects the majority base. The reconstructed nucleotide sequence is converted back into bytes, decompressed, and checked against the original SHA-256 fingerprint.")
flow(["Copy 0:  A C G T", "Copy 1:  A T G T   ← substitution", "Copy 2:  A C G T", "             ↑", "Consensus: A C G T   ✓"])
callout("Interpretation", "Redundancy corrects isolated substitutions when a majority of copies remains correct. SHA-256 detects silent corruption that redundancy cannot repair.")

doc.add_heading("4. Intelligent tier placement", level=1)
doc.add_paragraph("Every artifact receives an explainable placement score. Higher scores indicate that the artifact deserves faster storage. The MVP combines normalized access frequency, exponential recency, predicted demand, business priority, and a size penalty.")
flow(["score = 0.28·frequency + 0.22·recency + 0.28·demand", "+ 0.22·importance − 0.10·size penalty"])
fixed_table(["Score", "Recommended tier", "Meaning"], [
    ("0.72–1.00", "GPU", "Highest expected near-term value"),
    ("0.56–0.71", "RAM", "Frequently or soon required"),
    ("0.38–0.55", "SSD", "Moderate value; keep prepared"),
    ("0.20–0.37", "S3", "Rarely required"),
    ("Below 0.20", "DNA", "Long-term archival candidate"),
], [1800, 2200, 5360])
doc.add_paragraph("The controller is deterministic. This keeps storage movement auditable and avoids using an LLM for work that ordinary systems code can perform reliably.")

doc.add_heading("5. Predictive prefetching", level=1)
doc.add_paragraph("When a request arrives, the resolver extracts meaningful terms and ranks artifact IDs that contain those terms. For the request ‘Compare address agents using the 2024 evaluation dataset,’ it identifies artifacts such as address-agent-2024 and evaluation-dataset-2024.")
fixed_table(["Without prefetch", "With prefetch"], [
    ("Plan request", "Planning and retrieval begin together"),
    ("Discover missing artifact", "Cold artifact is already being restored"),
    ("Wait for DNA retrieval", "Wait only for the slower parallel branch"),
    ("Load and infer", "Load and infer"),
], [4680, 4680])
doc.add_heading("Benchmark model", level=2)
doc.add_paragraph("The portfolio benchmark is a transparent simulation, not a wall-clock measurement of physical DNA. It uses 1,200 ms for planning, tier-dependent retrieval latency, 100 ms for loading, and 500 ms for inference. With DNA as the slowest dependency, the demonstrated result is:")
fixed_table(["Scenario", "Modeled latency", "Result"], [
    ("Without prefetch", "4,300 ms", "Planning followed by retrieval"),
    ("With prefetch", "3,100 ms", "Planning overlaps retrieval"),
    ("Difference", "1,200 ms", "27.9% improvement"),
], [3000, 2600, 3760])

doc.add_heading("6. Practical validation", level=1)
doc.add_paragraph("The project contains five automated tests and a browser-guided workflow. Together they validate the following behaviors:")
for item in [
    "Arbitrary bytes survive DNA encoding, FASTA serialization, parsing, decoding, and decompression.",
    "Replicated strands recover the original data after independent nucleotide substitutions.",
    "A cold artifact is resolved from a request, restored from DNA, and prefetched to SSD.",
    "A real 4 KB binary fixture survives DNA archival, 20 injected mutations, and verified restoration.",
    "The benchmark reports a lower modeled latency when retrieval overlaps planning."
]: bullet(item)
callout("Verified status", "All five automated tests pass. A live smoke test also recovered a real uploaded payload after 12 mutations and downloaded the exact original content.")

doc.add_heading("7. What is real and what is simulated", level=1)
fixed_table(["Implemented as real software", "Represented as a simulation"], [
    ("Compression and decompression", "Physical DNA synthesis and sequencing"),
    ("Binary-to-nucleotide encoding", "Actual GPU and RAM allocation"),
    ("FASTA-like archive files", "Cloud object storage"),
    ("Substitution injection and repair", "Production retrieval latency and cost"),
    ("SHA-256 integrity checks", "Model inference"),
    ("File upload and byte-exact download", "Large-scale multi-node orchestration"),
    ("Placement and prefetch decisions", "LLM/embedding-based request understanding"),
], [4680, 4680])
doc.add_paragraph("This distinction is essential. The project researches the computer-systems layer around DNA archival storage; it does not claim that a local .dna text file reproduces molecular synthesis, preservation, random access, or sequencing.")

doc.add_heading("8. Limitations", level=1)
for item in [
    "The two-bit codec can create long homopolymers and does not actively control GC balance.",
    "Triple replication handles limited substitutions but is less efficient and capable than Reed–Solomon or fountain-code designs.",
    "The resolver uses filename keyword overlap rather than semantic embeddings or an LLM planner.",
    "Tier directories represent infrastructure abstractions; there is no CUDA loader, Redis cache, database, or S3 provider.",
    "Benchmark values are modeled constants and should not be presented as physical DNA performance measurements.",
    "The registry is a JSON file and is not designed for concurrent production workloads."
]: bullet(item)

doc.add_heading("9. Future plan", level=1)
doc.add_heading("Phase 2 — biologically aware codec", level=2)
for item in ["Add Reed–Solomon error correction and missing-strand recovery.", "Constrain homopolymer length and maintain GC balance.", "Model substitution, insertion, deletion, and strand-loss errors.", "Embed strand indexes and checksums in the encoded sequence."]: bullet(item)
doc.add_heading("Phase 3 — real infrastructure", level=2)
for item in ["Replace tier folders with Redis/RAM cache, local or NVMe SSD, and real S3-compatible object storage.", "Load a small LoRA adapter through a genuine inference runtime.", "Move registry and event history into SQLite or PostgreSQL.", "Collect wall-clock latency, cache-hit rate, storage cost, and prefetch waste."]: bullet(item)
doc.add_heading("Phase 4 — smarter agentic control", level=2)
for item in ["Use embeddings or an LLM only for dependency prediction and planning.", "Train or evaluate a demand-forecasting model from access history.", "Support multi-artifact plans and cancellation of incorrect prefetches.", "Compare rule-based, learned, and hybrid placement policies on the same workload."]: bullet(item)

doc.add_heading("10. Portfolio positioning", level=1)
doc.add_paragraph("HelixCache is best presented as an AI systems and storage project, not merely a DNA encoder. Its distinctive contribution is the combination of hierarchical caching, explainable placement, archival error recovery, and prefetch scheduling around an agent’s future needs.")
callout("One-sentence description", "HelixCache is an autonomous storage controller that predicts AI artifact demand and dynamically tiers models and knowledge across GPU, RAM, SSD, object storage, and simulated DNA storage—with verified recovery and predictive prefetching.")
doc.add_heading("Recommended next milestone", level=2)
doc.add_paragraph("The strongest next milestone is a measured end-to-end experiment using a real small adapter, real object storage, and an embedding-based resolver. Report cache-hit rate, prediction precision, latency saved, storage cost, and the penalty of incorrect prefetches across a reproducible request trace.")

doc.add_heading("Appendix A — Repository map", level=1)
fixed_table(["Path", "Purpose"], [
    ("src/dna-codec.js", "DNA conversion, FASTA handling, mutation, consensus recovery, and archive analysis."),
    ("src/helix-cache.js", "Registry, placement policy, movement, retrieval, experiments, request resolution, and benchmark."),
    ("src/server.js", "Zero-dependency HTTP server and API routing."),
    ("public/", "Interactive portfolio dashboard."),
    ("test/", "Automated codec, recovery, artifact, prefetch, and benchmark tests."),
    ("data/", "Runtime registry and logical tier directories; ignored by Git."),
], [2600, 6760])

doc.core_properties.title = "HelixCache — Technical Project Report"
doc.core_properties.subject = "Agentic hierarchical memory using simulated DNA archival storage"
doc.core_properties.keywords = "AI systems, DNA storage, caching, prefetching, error correction"
doc.core_properties.author = "HelixCache Project"
doc.save(OUT)
print(OUT)
